import os
import math
import pandas as pd
from difflib import SequenceMatcher
import itertools

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import tkinter as tk
from tkinter import filedialog, messagebox


# ==============================
# Класс для поиска фото по ФИО и формирования отчёта
# ==============================
class FioPhotoMatcher:
    def __init__(self, excel_file, photo_folder, output_file="result.docx", verbose=False,
                 columns=5, photo_width_cm=2.5, photo_height_cm=3.5):
        """
        Инициализация объекта.

        :param excel_file: путь к файлу Excel с ФИО (ожидается колонка "ФИО")
        :param photo_folder: папка с фотографиями
        :param output_file: имя выходного Word-документа
        :param verbose: флаг вывода логов в консоль
        :param columns: количество колонок в таблице с фотографиями
        :param photo_width_cm: ширина фото в сантиметрах
        :param photo_height_cm: высота фото в сантиметрах
        """
        self.excel_file = excel_file
        self.photo_folder = photo_folder
        self.output_file = output_file
        self.verbose = verbose
        self.columns = columns
        self.photo_width_cm = photo_width_cm
        self.photo_height_cm = photo_height_cm

        self.names_list = []  # Список ФИО из Excel
        self.matched_list = []  # [(номер, ФИО, имя файла, метод поиска, схожесть, использованный вариант)]
        self.unmatched_list = []  # [(номер, ФИО)]

        # Словарь для замены символов
        self.char_mapping = {
            "Қ": "К", "қ": "к",
            "Ұ": "У", "ұ": "у",
            "Ә": "А", "ә": "а",
            "Ң": "Н", "ң": "н",
            "Ғ": "Г", "ғ": "г",
            "І": "Ы", "і": "ы",
            "Ө": "О", "ө": "о",
            "Ү": "У", "ү": "у"
        }

    def log(self, message):
        """Вывод сообщения в консоль, если включён режим VERBOSE."""
        if self.verbose:
            print(message)

    @staticmethod
    def set_columns(section, num_columns):
        """
        Устанавливает число столбцов в разделе section.
        """
        sectPr = section._sectPr
        cols = sectPr.find(qn('w:cols'))
        if cols is None:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        cols.set(qn('w:num'), str(num_columns))

    def read_fio_from_excel(self):
        """Читает ФИО из Excel (ожидается колонка 'ФИО') и сохраняет в self.names_list."""
        try:
            df = pd.read_excel(self.excel_file)
            if 'ФИО' in df.columns:
                self.names_list = df['ФИО'].dropna().apply(lambda x: str(x).strip()).tolist()
            else:
                self.log("В файле нет колонки 'ФИО'.")
        except FileNotFoundError:
            self.log(f"Файл '{self.excel_file}' не найден.")

    @staticmethod
    def compute_similarity(a, b):
        """Вычисляет процентную схожесть двух строк (без учета регистра)."""
        return SequenceMatcher(None, a.lower(), b.lower()).ratio() * 100

    @staticmethod
    def remove_extension(filename):
        """Возвращает имя файла без расширения."""
        return os.path.splitext(filename)[0]

    @staticmethod
    def shorten_fio(fio):
        """Возвращает укороченную версию ФИО (Фамилия + Имя)."""
        parts = fio.split()
        if len(parts) >= 2:
            return parts[0] + " " + parts[1]
        return fio

    @staticmethod
    def generate_variants_combinations(s, mapping, max_subs=7):
        """
        Генерирует варианты строки s, заменяя символы согласно mapping,
        выполняя от 1 до max_subs замен на различных позициях.
        Возвращает список уникальных вариантов.
        """
        indices = [i for i, ch in enumerate(s) if ch in mapping]
        variants = set()
        for r in range(1, max_subs + 1):
            for combo in itertools.combinations(indices, r):
                s_list = list(s)
                for idx in combo:
                    s_list[idx] = mapping[s_list[idx]]
                variant = "".join(s_list)
                variants.add(variant)
        return list(variants)

    def iterative_match(self, fio, filename, max_subs=7):
        """
        Пошагово сравнивает ФИО с именем файла (без расширения):
          1. Точное совпадение полного ФИО.
          2. Генерация вариантов полного ФИО с заменами.
          3. Точное совпадение укороченного ФИО (Фамилия + Имя).
          4. Генерация вариантов укороченного ФИО.
        Возвращает кортеж: (найдено: bool, метод: str, схожесть: float, использованный вариант: str)
        """
        file_base = self.remove_extension(filename)

        # 1. Точное совпадение полного ФИО
        if fio.lower() in file_base.lower():
            sim = self.compute_similarity(fio, file_base)
            return True, "Full FIO exact", sim, fio

        # 2. Генерация вариантов для полного ФИО
        variants = self.generate_variants_combinations(fio, self.char_mapping, max_subs)
        for variant in variants:
            if variant.lower() in file_base.lower():
                sim = self.compute_similarity(variant, file_base)
                return True, "Full FIO substitution (combination)", sim, variant

        # 3. Точное совпадение укороченного ФИО
        fio_short = self.shorten_fio(fio)
        if fio_short.lower() in file_base.lower():
            sim = self.compute_similarity(fio_short, file_base)
            return True, "Short FIO exact", sim, fio_short

        # 4. Генерация вариантов для укороченного ФИО
        variants_short = self.generate_variants_combinations(fio_short, self.char_mapping, max_subs)
        for variant in variants_short:
            if variant.lower() in file_base.lower():
                sim = self.compute_similarity(variant, file_base)
                return True, "Short FIO substitution (combination)", sim, variant

        return False, "No match", 0, ""

    def search_photo_for_fio_iterative(self, fio, max_subs=7):
        """
        Перебирает файлы в папке self.photo_folder и возвращает первый файл,
        удовлетворяющий логике итеративного поиска ФИО.
        Возвращает (имя файла, метод, схожесть, использованный вариант).
        Если не найдено, возвращает (None, "No match", 0, "").
        """
        try:
            files = os.listdir(self.photo_folder)
        except FileNotFoundError:
            self.log(f"Папка '{self.photo_folder}' не найдена.")
            return None, "Folder not found", 0, ""

        for file in files:
            if not file.lower().endswith(('.jpg', '.jpeg', '.png')):
                continue
            found, method, sim, variant_used = self.iterative_match(fio, file, max_subs)
            if found:
                return file, method, sim, variant_used
        return None, "No match", 0, ""

    def process(self):
        """Основной метод: считывает ФИО, ищет фотографии, формирует отчёт."""
        self.read_fio_from_excel()
        if not self.names_list:
            self.log("Список ФИО пуст.")
            return

        total_count = len(self.names_list)
        match_count = 0

        for i, fio in enumerate(self.names_list, start=1):
            file, method, sim, variant_used = self.search_photo_for_fio_iterative(fio)
            if file:
                match_count += 1
                self.log(
                    f"{i}. ФИО: '{fio}' -> файл: '{file}' | Метод: {method} | Схожесть: {sim:.1f}% | Вариант: '{variant_used}'")
                self.matched_list.append((i, fio, file, method, sim, variant_used))
            else:
                self.log(f"{i}. ФИО: '{fio}' -> Файл не найден.")
                self.unmatched_list.append((i, fio))

        self.log(f"\nОбщее количество ФИО: {total_count}")
        self.log(f"Количество совпадений: {match_count}")

        self.create_word_report()

    def create_word_report(self):
        """
        Создает Word-документ с таблицей фотографий и разделом отчёта,
        оформленным в два столбца.
        """
        doc = Document()

        # Таблица с фотографиями
        photo_width = Cm(self.photo_width_cm)
        photo_height = Cm(self.photo_height_cm)
        from math import ceil
        row_count = ceil(len(self.matched_list) / self.columns)
        table = doc.add_table(rows=row_count, cols=self.columns)
        table.style = 'Table Grid'

        for idx, (num, fio, file, method, sim, variant_used) in enumerate(self.matched_list):
            row_idx = idx // self.columns
            col_idx = idx % self.columns
            cell = table.cell(row_idx, col_idx)
            photo_path = os.path.join(self.photo_folder, file)
            if os.path.exists(photo_path):
                run = cell.paragraphs[0].add_run()
                try:
                    run.add_picture(photo_path, width=photo_width, height=photo_height)
                except Exception as e:
                    self.log(f"Ошибка при вставке фото {photo_path}: {e}")
            p = cell.add_paragraph()
            run_text = p.add_run(f"{num}) {fio}\nСхожесть: {sim:.1f}%")
            run_text.font.size = Pt(7)
            run_text.font.name = 'Arial'

        # Новый раздел для отчёта (с двухколоночным макетом)
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        self.set_columns(new_section, 2)

        doc.add_heading("Отчёт по результатам", level=1)

        p_total = doc.add_paragraph()
        r_total = p_total.add_run(f"Всего ФИО: {len(self.matched_list) + len(self.unmatched_list)}")
        r_total.font.size = Pt(7)

        p_found = doc.add_paragraph()
        r_found = p_found.add_run(f"Найдено совпадений: {len(self.matched_list)}")
        r_found.font.size = Pt(7)

        p_not_found = doc.add_paragraph()
        r_not_found = p_not_found.add_run(f"Не найдено совпадений: {len(self.unmatched_list)}")
        r_not_found.font.size = Pt(7)

        # Список найденных ФИО
        if self.matched_list:
            p_list_found = doc.add_paragraph("Список найденных ФИО:", style='List Bullet')
            if p_list_found.runs:
                p_list_found.runs[0].font.size = Pt(7)
            for (num, fio, file, method, sim, variant_used) in self.matched_list:
                p_item = doc.add_paragraph(f"{num}) {fio}", style='List Bullet 2')
                if p_item.runs:
                    p_item.runs[0].font.size = Pt(7)

        # Список не найденных ФИО
        if self.unmatched_list:
            p_list_not_found = doc.add_paragraph("Список не найденных ФИО:", style='List Bullet')
            if p_list_not_found.runs:
                p_list_not_found.runs[0].font.size = Pt(7)
            for (num, fio) in self.unmatched_list:
                p_item = doc.add_paragraph(f"{num}) {fio}", style='List Bullet 2')
                if p_item.runs:
                    p_item.runs[0].font.size = Pt(7)

        try:
            doc.save(self.output_file)
            self.log(f"Документ сохранён как {self.output_file}")
        except Exception as e:
            self.log(f"Ошибка сохранения документа: {e}")


# ==============================
# Графический интерфейс на базе tkinter
# ==============================
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Поиск фотографий по ФИО")
        self.geometry("700x320")
        self.resizable(False, False)

        # Путь к Excel-файлу
        tk.Label(self, text="Путь к Excel-файлу:").grid(row=0, column=0, sticky="w", padx=5, pady=5)
        self.entry_excel = tk.Entry(self, width=50)
        self.entry_excel.grid(row=0, column=1, padx=5, pady=5)
        tk.Button(self, text="Обзор", command=self.browse_excel).grid(row=0, column=2, padx=5, pady=5)

        # Путь к папке с фотографиями
        tk.Label(self, text="Путь к папке с фотографиями:").grid(row=1, column=0, sticky="w", padx=5, pady=5)
        self.entry_photo_folder = tk.Entry(self, width=50)
        self.entry_photo_folder.grid(row=1, column=1, padx=5, pady=5)
        tk.Button(self, text="Обзор", command=self.browse_photo_folder).grid(row=1, column=2, padx=5, pady=5)

        # Путь для сохранения Word-документа
        tk.Label(self, text="Путь для сохранения Word-документа:").grid(row=2, column=0, sticky="w", padx=5, pady=5)
        self.entry_output = tk.Entry(self, width=50)
        self.entry_output.grid(row=2, column=1, padx=5, pady=5)
        tk.Button(self, text="Обзор", command=self.browse_output).grid(row=2, column=2, padx=5, pady=5)

        # Размеры фотографий
        tk.Label(self, text="Ширина фото (см):").grid(row=3, column=0, sticky="w", padx=5, pady=5)
        self.entry_width = tk.Entry(self, width=10)
        self.entry_width.insert(0, "2.5")
        self.entry_width.grid(row=3, column=1, sticky="w", padx=5, pady=5)

        tk.Label(self, text="Высота фото (см):").grid(row=4, column=0, sticky="w", padx=5, pady=5)
        self.entry_height = tk.Entry(self, width=10)
        self.entry_height.insert(0, "3.5")
        self.entry_height.grid(row=4, column=1, sticky="w", padx=5, pady=5)

        # Количество колонок в таблице с фотографиями
        tk.Label(self, text="Количество колонок:").grid(row=5, column=0, sticky="w", padx=5, pady=5)
        self.entry_columns = tk.Entry(self, width=10)
        self.entry_columns.insert(0, "5")
        self.entry_columns.grid(row=5, column=1, sticky="w", padx=5, pady=5)

        # Кнопка запуска
        tk.Button(self, text="Запустить", command=self.start_processing, width=20).grid(row=6, column=1, pady=15)

    def browse_excel(self):
        file_path = filedialog.askopenfilename(title="Выберите Excel-файл", filetypes=[("Excel files", "*.xlsx *.xls")])
        if file_path:
            self.entry_excel.delete(0, tk.END)
            self.entry_excel.insert(0, file_path)

    def browse_photo_folder(self):
        folder_path = filedialog.askdirectory(title="Выберите папку с фотографиями")
        if folder_path:
            self.entry_photo_folder.delete(0, tk.END)
            self.entry_photo_folder.insert(0, folder_path)

    def browse_output(self):
        file_path = filedialog.asksaveasfilename(title="Сохранить Word-документ", defaultextension=".docx",
                                                 filetypes=[("Word Document", "*.docx")])
        if file_path:
            self.entry_output.delete(0, tk.END)
            self.entry_output.insert(0, file_path)

    def start_processing(self):
        excel_file = self.entry_excel.get()
        photo_folder = self.entry_photo_folder.get()
        output_file = self.entry_output.get()

        try:
            photo_width_cm = float(self.entry_width.get())
            photo_height_cm = float(self.entry_height.get())
        except ValueError:
            messagebox.showerror("Ошибка", "Размер фото должен быть числовым значением.")
            return

        try:
            columns = int(self.entry_columns.get())
        except ValueError:
            messagebox.showerror("Ошибка", "Количество колонок должно быть целым числом.")
            return

        if not os.path.exists(excel_file):
            messagebox.showerror("Ошибка", "Excel-файл не найден.")
            return

        if not os.path.isdir(photo_folder):
            messagebox.showerror("Ошибка", "Папка с фотографиями не найдена.")
            return

        matcher = FioPhotoMatcher(
            excel_file=excel_file,
            photo_folder=photo_folder,
            output_file=output_file,
            verbose=False,  # Можно изменить на True для вывода логов в консоль
            columns=columns,
            photo_width_cm=photo_width_cm,
            photo_height_cm=photo_height_cm
        )
        matcher.process()
        messagebox.showinfo("Готово", f"Документ сохранён как:\n{output_file}")


if __name__ == "__main__":
    app = App()
    app.mainloop()

